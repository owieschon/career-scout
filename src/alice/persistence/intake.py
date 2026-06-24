"""intake — orchestrate multi-modal onboarding into a confirmable Profile.

Flow:  resume file / voice note / chat text
         -> parse_resume / transcribe_voice (raw text)
         -> profile_gen.extract (grounded fields)
         -> profile_gen.build_profile (+ optional archetype seed)
         -> profile_store.save_draft (UNCONFIRMED)
         -> build_confirm_payload (shown to user)
         -> [user confirms] -> confirm_intake -> profile_store.confirm

CONFIRM-THEN-COMMIT IS THE WHOLE POINT. run_intake never produces a usable
profile; it produces a DRAFT plus the text to show the user. Only confirm_intake
flips the grounding gate. Nothing downstream can source or score until then
(profile_store.load_active returns None for a draft).

Resume parsing: .docx (python-docx), .pdf (PyMuPDF/fitz), .txt/.md (raw).
Voice: AssemblyAI REST (US region, no-train) — see transcribe_voice. The HTTP
client is urllib (matches llm.py); no SDK dependency added. The provider call is
injectable so tests run offline.
"""
from __future__ import annotations

import json
import ssl
import time
import urllib.request
from pathlib import Path

from alice.persistence import profile_gen
from alice.persistence import profile_store
from alice.persistence.profile_schema import Profile

try:
    import certifi
    _SSL = ssl.create_default_context(cafile=certifi.where())
except ImportError:
    _SSL = ssl.create_default_context()


class IntakeError(Exception):
    """Raised for unrecoverable intake problems (unsupported file, missing voice
    provider key, etc.). Fail loud (P2) — never degrade to a fabricated profile."""


# ── resume parsing ────────────────────────────────────────────────────────────
def parse_resume(path: str | Path) -> str:
    """Extract plain text from a resume file. Supports .docx / .pdf / .txt / .md.

    Raises IntakeError on an unsupported extension or empty extraction — a
    resume we can't read must fail loud, not silently yield an empty profile.
    """
    path = Path(path)
    if not path.exists():
        raise IntakeError(f"resume file not found: {path}")
    ext = path.suffix.lower()

    if ext == ".docx":
        text = _parse_docx(path)
    elif ext == ".pdf":
        text = _parse_pdf(path)
    elif ext in (".txt", ".md"):
        text = path.read_text(errors="replace")
    else:
        raise IntakeError(
            f"unsupported resume type {ext!r}; send a .docx, .pdf, .txt, or .md"
        )

    text = (text or "").strip()
    if not text:
        raise IntakeError(f"could not extract any text from {path.name}")
    return text


def _parse_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
 # tables carry real content in many resume templates
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _parse_pdf(path: Path) -> str:
    import fitz  # PyMuPDF
    parts = []
    with fitz.open(str(path)) as doc:
        for page in doc:
            parts.append(page.get_text())
    return "\n".join(parts)


# ── voice transcription (AssemblyAI, US region, no-train) ──────────────────────
# AssemblyAI's default host (api.assemblyai.com) is US-hosted and does not train
# on customer audio (no opt-in here). EU residency would use api.eu.assemblyai.com
# — we intentionally pin the US host.
_AAI_BASE = "https://api.assemblyai.com/v2"


def transcribe_voice(path: str | Path, cfg: dict | None = None,
                     *, transcriber=None, poll_interval: float = 2.0,
                     timeout_s: float = 180.0) -> str:
    """Transcribe a voice note to text via a US, no-train provider.

    `transcriber` is injectable (a callable taking the file path, returning
    text) so tests and alternate providers don't touch the network. The default
    path uses AssemblyAI and requires ASSEMBLYAI_API_KEY in config.env; without
    it, fail loud with an actionable message rather than guessing a transcript.
    """
    path = Path(path)
    if not path.exists():
        raise IntakeError(f"voice file not found: {path}")

    if transcriber is not None:
        text = transcriber(str(path))
        if not (text or "").strip():
            raise IntakeError("voice transcription returned empty text")
        return text.strip()

    if cfg is None:
        from alice import jobcfg
        cfg = jobcfg.load()
    key = cfg.get("ASSEMBLYAI_API_KEY")
    if not key:
        raise IntakeError(
            "voice intake needs ASSEMBLYAI_API_KEY in ~/.config/job-search/config.env "
            "(AssemblyAI, US-hosted, no-train). Add the key, or send your resume as a "
            "document and tell me the rest in chat."
        )
    return _assemblyai_transcribe(path, key, poll_interval=poll_interval, timeout_s=timeout_s)


def _aai_request(url: str, key: str, *, data: bytes | None = None,
                 json_body: dict | None = None, method: str = "GET") -> dict:
    headers = {"authorization": key}
    body = data
    if json_body is not None:
        body = json.dumps(json_body).encode()
        headers["content-type"] = "application/json"
    req = urllib.request.Request(url, data=body, headers=headers, method=method)
    with urllib.request.urlopen(req, context=_SSL, timeout=60) as resp:
        return json.loads(resp.read().decode("utf-8", "replace"))


def _assemblyai_transcribe(path: Path, key: str, *, poll_interval: float,
                           timeout_s: float) -> str:
 # 1. upload audio
    up = _aai_request(f"{_AAI_BASE}/upload", key, data=path.read_bytes(), method="POST")
    audio_url = up.get("upload_url")
    if not audio_url:
        raise IntakeError("AssemblyAI upload failed (no upload_url returned)")

 # 2. request transcription
    job = _aai_request(
        f"{_AAI_BASE}/transcript", key,
        json_body={"audio_url": audio_url}, method="POST",
    )
    tid = job.get("id")
    if not tid:
        raise IntakeError("AssemblyAI transcript request failed (no id returned)")

 # 3. poll
    deadline = time.monotonic() + timeout_s
    while time.monotonic() < deadline:
        st = _aai_request(f"{_AAI_BASE}/transcript/{tid}", key)
        status = st.get("status")
        if status == "completed":
            text = (st.get("text") or "").strip()
            if not text:
                raise IntakeError("AssemblyAI returned an empty transcript")
            return text
        if status == "error":
            raise IntakeError(f"AssemblyAI transcription error: {st.get('error')}")
        time.sleep(poll_interval)
    raise IntakeError("AssemblyAI transcription timed out")


# ── orchestration ──────────────────────────────────────────────────────────────
def run_intake(
    user_id: str,
    *,
    resume_path: str | Path | None = None,
    resume_text: str | None = None,
    chat_text: str | None = None,
    voice_path: str | Path | None = None,
    voice_text: str | None = None,
    archetype_key: str | None = None,
    identity_overrides: dict | None = None,
    llm_call=None,
    transcriber=None,
    cfg: dict | None = None,
) -> dict:
    """Run intake across whatever modalities are provided. Produces a DRAFT
    profile and the confirm payload. Does NOT commit.

    Returns {"profile": Profile (unconfirmed), "confirm_text": str,
             "sources": {kind: text}}. The caller shows confirm_text and, on the
    user's explicit yes, calls confirm_intake(user_id).
    """
    sources: dict[str, str] = {}

    if resume_text is not None:
        sources["resume"] = resume_text
    elif resume_path is not None:
        sources["resume"] = parse_resume(resume_path)

    if voice_text is not None:
        sources["voice"] = voice_text
    elif voice_path is not None:
        sources["voice"] = transcribe_voice(voice_path, cfg, transcriber=transcriber)

    if chat_text:
        sources["chat"] = chat_text

    if not sources and not archetype_key:
        raise IntakeError(
            "nothing to onboard from: provide a resume, a voice note, some chat "
            "text, or pick an archetype to start from"
        )

 # Extraction order = trust order (resume < chat < voice handled in build_profile
 # by application order; later overrides earlier). We extract resume, chat, voice.
    extractions = []
    for kind in ("resume", "chat", "voice"):
        if kind in sources:
            extractions.append(profile_gen.extract(sources[kind], kind, llm_call=llm_call))

    profile = profile_gen.build_profile(
        user_id,
        extractions,
        archetype_key=archetype_key,
        identity_overrides=identity_overrides,
    )

 # Belt-and-suspenders grounding check before we even store the draft. If the
 # audit ever finds an ungrounded "grounded" claim, fail loud — that's a bug
 # in extraction, and committing it would violate Rule 1.
    violations = profile_gen.no_fabrication_audit(profile, sources)
    if violations:
        raise IntakeError(
            "refusing to build a profile: grounding audit found unsupported "
            "claims:\n  - " + "\n  - ".join(violations)
        )

    profile_store.save_draft(profile)
 # Persist the raw sources so post-confirm steps (resume-variant derivation)
 # can re-read the originals without a re-upload.
    profile_store.save_sources(user_id, sources)
    confirm_text = profile_gen.build_confirm_payload(profile)
    return {"profile": profile, "confirm_text": confirm_text, "sources": sources}


def derive_variants_for(user_id: str, *, llm_call=None) -> dict:
    """Auto-derive track-tailored resume variants from the user's uploaded
    resume. Intended to run AFTER confirmation (so we never spend tokens on an
    unconfirmed profile). Returns {track: variant_dict}, or {} if the user never
    uploaded a resume (chat/voice-only onboarding). Never raises on a missing
    resume — variants are a bonus, not a gate."""
    from alice.persistence import generate_resume_variants

    sources = profile_store.load_sources(user_id)
    resume_text = sources.get("resume")
    if not resume_text:
        return {}
    return generate_resume_variants.derive_variants(
        resume_text, user_id, write=True, llm_call=llm_call,
    )


def confirm_intake(user_id: str) -> Profile:
    """Commit half of confirm-then-commit: flip the draft to confirmed. After
    this, profile_store.load_active(user_id) returns the profile and the engine
    may source/score on it. Raises if there's no draft to confirm."""
    return profile_store.confirm(user_id)


def cancel_intake(user_id: str) -> bool:
    """Discard the pending draft (confirm-cancel)."""
    return profile_store.discard_draft(user_id)
